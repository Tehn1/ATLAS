# !/usr/bin/python3
# -*- coding: utf-8 -*-


# AuTomated LAbbook System (ATLAS)

import sys
import os
import docx
import sqlite3

from PyQt5.QtWidgets import QMainWindow, QApplication, QDesktopWidget, \
    QFrame, QLabel, QPushButton, QCheckBox, QAction, QFileDialog, QTableView, \
    QAbstractItemView, QTabWidget, QWidget, QLineEdit, QPlainTextEdit, \
    QCalendarWidget, QDateEdit, QComboBox

from PyQt5.QtGui import QFont, QStandardItemModel, QStandardItem, QPalette, QColor

#Key Global Variables

Experiments = {}

#Calendar for Date Selection
class Cal(QMainWindow):
    def __init__(self, parent=None):
        super(Cal, self).__init__(parent)
        self.setWindowTitle("Choose Section Date")
        self.resize(310,210)
        self.move(750,250)

        self.Cale = QCalendarWidget(self)
        self.Cale.setVerticalHeaderFormat(self.Cale.NoVerticalHeader)
        self.Cale.move(5,5)
        self.Cale.resize(300,200)

        self.Cale.clicked()

#Main Window
class Atlas(QMainWindow):

    def __init__(self):
        super().__init__()

        self.initUI()

    def initUI(self):

        #On startup, check whether variable databases exist. If not, build databases.

        #Check Variable Database
        if os.path.exists("./Variables/Variables.sqlite3") == False:
            #Create Variables Database
            conn = sqlite3.connect("./Variables/Variables.sqlite3")
            c = conn.cursor()

            #Create Antibiotics Table
            c.execute("CREATE TABLE Antibiotics(id INTEGER PRIMARY KEY, name TEXT)")
            ant="Carbenicillin"
            c.execute("INSERT INTO Antibiotics(name) VALUES (?)",(ant,))
            ant="Kanamycin"
            c.execute("INSERT INTO Antibiotics(name) VALUES (?)",(ant,))
            ant="Chloramphenicol"
            c.execute("INSERT INTO Antibiotics(name) VALUES (?)",(ant,))

            # Create Plasmids Table
            c.execute("CREATE TABLE Plasmids(id INTEGER PRIMARY KEY, name TEXT)")

            # Create Primers Table
            c.execute("CREATE TABLE Primers(id INTEGER PRIMARY KEY, name TEXT)")

            conn.commit()
            conn.close()

        #Insert Menu Bar
        self.openFile = QAction('&Open', self)
        self.openFile.setShortcut('Ctrl+O')
        self.openFile.setStatusTip('Open new File')
        self.openFile.triggered.connect(self.showDialog)

        menubar = self.menuBar()
        self.fileMenu = menubar.addMenu('&File')
        self.fileMenu.addAction(self.openFile)

        #Center Window
        self.resize(1200, 600)
        qtRectangle = self.frameGeometry()
        centerPoint = QDesktopWidget().availableGeometry().center()
        qtRectangle.moveCenter(centerPoint)
        self.move(qtRectangle.topLeft())
        self.setWindowTitle('ATLAS')

        #Insert Status Bar
        self.statusBar().showMessage('Ready')
        self.show()

        #Experiment Selection Frame
        self.ExpFrame = QFrame(self)
        self.ExpFrame.move(5, 25)
        self.ExpFrame.resize(450, 200)
        self.ExpFrame.setFrameShape(QFrame.StyledPanel)
        self.ExpFrame.show()

        #Experiment Frame Label
        self.ExpLabel = QLabel(self.ExpFrame)
        self.ExpLabel.setText("Experiment Tables")
        self.ExpLabel.move(5, 1)
        newfont = QFont("Times", 8, QFont.Bold)
        self.ExpLabel.setFont(newfont)
        self.ExpLabel.show()

        #Experiment Table Generation Button
        self.ExpButton = QPushButton(self.ExpFrame)
        self.ExpButton.resize(120, 30)
        self.ExpButton.move(320, 20)
        self.ExpButton.setText("Generate")
        self.ExpButton.clicked.connect(self.GenExpList)
        self.ExpButton.show()

        #Experiment Check Buttons
        self.ExpCheck1 = QCheckBox(self.ExpFrame)
        self.ExpCheck1.move(15, 30)
        self.ExpCheck1.setText("Include Complete Experiments")
        self.ExpCheck1.show()

        #Experiment Table
        self.ExpTable = QTableView(self.ExpFrame)
        self.ExpTable.move(10, 60)
        self.ExpTable.resize(430, 130)
        self.ExpTableModel = QStandardItemModel(self)
        self.ExpTable.setModel(self.ExpTableModel)
        self.ExpTable.setEditTriggers(QAbstractItemView.NoEditTriggers)
        row = []
        cell = QStandardItem("#")
        row.append(cell)
        cell = QStandardItem("Name")
        row.append(cell)
        self.ExpTable.horizontalHeader().hide()
        self.ExpTable.verticalHeader().hide()
        self.ExpTableModel.appendRow(row)
        self.ExpTable.horizontalHeader().setStretchLastSection(True)
        self.ExpTable.setColumnWidth(0,12)
        self.ExpTable.resizeRowsToContents()
        self.ExpTable.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.ExpTable.show()

        #Modification Frame
        self.ModFrame = QFrame(self)
        self.ModFrame.move(5, 230)
        self.ModFrame.resize(450, 350)
        self.ModFrame.setFrameShape(QFrame.StyledPanel)
        self.ModFrame.show()

        #Modification Label
        self.ModLabel = QLabel(self.ModFrame)
        self.ModLabel.setText("Section Viewer")
        self.ModLabel.move(5, 1)
        newfont = QFont("Times", 8, QFont.Bold)
        self.ModLabel.setFont(newfont)
        self.ModLabel.show()

        #Modification Table
        self.ModTable = QTableView(self.ModFrame)
        self.ModTable.move(10, 20)
        self.ModTable.resize(430, 320)
        self.ModTableModel = QStandardItemModel(self)
        self.ModTable.setModel(self.ModTableModel)
        self.ModTable.setEditTriggers(QAbstractItemView.NoEditTriggers)
        row = []
        cell = QStandardItem("Date")
        row.append(cell)
        cell = QStandardItem("Name")
        row.append(cell)
        self.ModTable.horizontalHeader().hide()
        self.ModTable.verticalHeader().hide()
        self.ModTableModel.appendRow(row)
        self.ModTable.horizontalHeader().setStretchLastSection(True)
        self.ModTable.setColumnWidth(0,80)
        self.ModTable.resizeRowsToContents()
        self.ModTable.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.ModTable.clicked.connect(self.openExp)
        self.ModTable.show()

        #Detailed Tabs
        self.DetTabs = QTabWidget(self)
        self.DetTab1 = QWidget(self)
        self.DetTab1.setAutoFillBackground(True)
        self.DetTab2 = QWidget(self)
        self.DetTab2.setAutoFillBackground(True)
        self.DetTab3 = QWidget(self)
        self.DetTab3.setAutoFillBackground(True)
        self.DetTabs.move(460, 25)
        self.DetTabs.resize(735, 556)
        self.DetTabs.addTab(self.DetTab1,"DetTab1")
        self.DetTabs.addTab(self.DetTab2,"DetTab2")
        self.DetTabs.addTab(self.DetTab3, "New Protocol")
        self.DetTabs.show()

        self.DetTabs.currentChanged.connect(self.TabChange)

        #New Protocol Tab Setup
        self.DTNew_Cat_Title = QLabel(self.DetTab3)
        self.DTNew_Cat_Title.setText("Protocol Location")
        self.DTNew_Cat_Title.move(5, 2)
        newfont = QFont("Times", 8, QFont.Bold)
        self.DTNew_Cat_Title.setFont(newfont)
        self.DTNew_Cat_Title.show()

        self.DTNew_CatText_Title = QLineEdit(self.DetTab3)
        self.DTNew_CatText_Title.setText("./Protocols/General/")
        self.DTNew_CatText_Title.move(3,22)
        self.DTNew_CatText_Title.resize(723,17)
        self.DTNew_CatText_Title.show()

        self.CatButton = QPushButton(self.DetTab3)
        self.CatButton.resize(120, 22)
        self.CatButton.move(2, 40)
        self.CatButton.setText("Change Location...")
        self.CatButton.clicked.connect(self.showDialogDir)

        self.CatButton.show()

        self.DTNew_Lab_Title = QLabel(self.DetTab3)
        self.DTNew_Lab_Title.setText("Title")
        self.DTNew_Lab_Title.move(5, 65)
        newfont = QFont("Times", 8, QFont.Bold)
        self.DTNew_Lab_Title.setFont(newfont)
        self.DTNew_Lab_Title.show()

        self.DTNew_Text_Title = QLineEdit(self.DetTab3)
        self.DTNew_Text_Title.setText("Protocol Title")
        self.DTNew_Text_Title.move(3,83)
        self.DTNew_Text_Title.resize(723,17)
        self.DTNew_Text_Title.show()

        self.DTNew_Lab_Title = QLabel(self.DetTab3)
        self.DTNew_Lab_Title.setText("Select Section")
        self.DTNew_Lab_Title.move(5, 105)
        newfont = QFont("Times", 8, QFont.Bold)
        self.DTNew_Lab_Title.setFont(newfont)
        self.DTNew_Lab_Title.show()

        self.DTNew_Text_Title = QComboBox(self.DetTab3)
        self.DTNew_Text_Title.move(3,123)
        self.DTNew_Text_Title.resize(723,17)
        self.DTNew_Text_Title.show()


        self.DTNew_Sec_Name = QLabel(self.DetTab3)
        self.DTNew_Sec_Name.setText("Section Name")
        self.DTNew_Sec_Name.move(5, 147)
        newfont = QFont("Times", 8, QFont.Bold)
        self.DTNew_Sec_Name.setFont(newfont)
        self.DTNew_Sec_Name.show()

        self.DTNew_Sec_Text = QLineEdit(self.DetTab3)
        self.DTNew_Sec_Text.setText("Section Name")
        self.DTNew_Sec_Text.move(3,168)
        self.DTNew_Sec_Text.resize(723,17)
        self.DTNew_Sec_Text.show()

        # self.DTNew_Lab_Date = QLabel(self.DetTab3)
        # self.DTNew_Lab_Date.setText("Date")
        # self.DTNew_Lab_Date.move(5, 107)
        # newfont = QFont("Times", 8, QFont.Bold)
        # self.DTNew_Lab_Date.setFont(newfont)
        # self.DTNew_Lab_Date.show()
        #
        # self.DTNew_But_Date = QDateEdit(self.DetTab3)
        # self.DTNew_But_Date.move(3,128)
        # self.DTNew_But_Date.resize(80,19)
        # self.DTNew_But_Date.setCalendarPopup(True)
        # self.DTNew_But_Date.show()

        self.DTNew_Lab_Section = QLabel(self.DetTab3)
        self.DTNew_Lab_Section.setText("Section Text")
        self.DTNew_Lab_Section.move(5, 194)
        newfont = QFont("Times", 8, QFont.Bold)
        self.DTNew_Lab_Section.setFont(newfont)
        self.DTNew_Lab_Section.show()

        self.DTNew_Text_Section = QPlainTextEdit(self.DetTab3)
        self.DTNew_Text_Section.appendPlainText("Section Text")
        self.DTNew_Text_Section.move(3,216)
        self.DTNew_Text_Section.resize(723,200)
        self.DTNew_Text_Section.show()

        self.DTNew_Var_Lab = QLabel(self.DetTab3)
        self.DTNew_Var_Lab.setText("Insert Variable")
        self.DTNew_Var_Lab.move(5, 426)
        newfont = QFont("Times", 8, QFont.Bold)
        self.DTNew_Var_Lab.setFont(newfont)
        self.DTNew_Var_Lab.show()

        self.DTNew_Var_Comb = QComboBox(self.DetTab3)
        self.DTNew_Var_Comb.move(3,446)
        self.DTNew_Var_Comb.resize(100,30)
        self.DTNew_Var_Comb.currentTextChanged.connect(self.VarChange)
        self.DTNew_Var_Comb.show()

        self.DTNew_Var_Lab2 = QLabel(self.DetTab3)
        self.DTNew_Var_Lab2.setText("Default Value")
        self.DTNew_Var_Lab2.move(130, 426)
        newfont = QFont("Times", 8, QFont.Bold)
        self.DTNew_Var_Lab2.setFont(newfont)
        self.DTNew_Var_Lab2.show()

        self.DTNew_Var_Comb2 = QComboBox(self.DetTab3)
        self.DTNew_Var_Comb2.move(130,446)
        self.DTNew_Var_Comb2.resize(100,30)
        self.DTNew_Var_Comb2.show()

        self.DTNew_VarIns_But = QPushButton(self.DetTab3)
        self.DTNew_VarIns_But.setText("Insert")
        self.DTNew_VarIns_But.move(257,446)
        self.DTNew_VarIns_But.resize(100,30)
        self.DTNew_VarIns_But.clicked.connect(self.InsVar)
        self.DTNew_VarIns_But.show()

        self.DTNew_AddSec_Button = QPushButton(self.DetTab3)
        self.DTNew_AddSec_Button.move(3,486)
        self.DTNew_AddSec_Button.resize(300,30)
        self.DTNew_AddSec_Button.setText("Add Section")
        self.DTNew_AddSec_Button.show()

    #Read Variable List when New Protocol tab is opened
    def TabChange(self, i):
        if i == 2:
            print("yes")
            conn = sqlite3.connect("./Variables/Variables.sqlite3")
            tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table';")
            templist = []
            for name in tables:
                templist.append(name[0])
            self.DTNew_Var_Comb.addItems(templist)

            conn.close()


    #Open file when experiment is clicked in list - temporary function
    def openExp(self,clickedIndex):
        Exp = clickedIndex.sibling(clickedIndex.row(),0).data()
        print(Exp)
        os.startfile(Experiments[Exp])

    #Test function for opening file dialog box - temporary function
    def showDialog(self):
        fname = QFileDialog.getOpenFileName(self, 'Open file', '.')
        print(fname[0])

    #Open dialog box and select a folder
    def showDialogDir(self):
        dname = str(QFileDialog.getExistingDirectory(self, 'Select Directory'))
        self.DTNew_CatText_Title.setText(dname)

    #Search Experiments folder and build list of Experiments present on computer
    def GenExpList(self):
        self.ExpTableModel.setRowCount(0)
        for root, dirs, files in os.walk(".\Experiments"):
            for file in files:
                if file.startswith("Experiment"):
                    if "." in file.split()[1]:
                        doc = docx.Document(root + "\\" + file)
                        if doc.core_properties.subject != "Complete" or self.ExpCheck1.isChecked():
                            row = []
                            row.append(QStandardItem(file.split()[1][:file.split()[1].index(".")]))
                            row.append(QStandardItem(doc.core_properties.title))
                            self.ExpTableModel.appendRow(row)
                            self.ExpTable.resizeRowsToContents()
                            Experiments[file.split()[1][:file.split()[1].index(".")]] = root+"\\"+file

                    else:
                        doc = docx.Document(root+"\\"+file)
                        if doc.core_properties.subject != "Complete" or self.ExpCheck1.isChecked():
                            row = []
                            row.append(QStandardItem(file.split()[1]))
                            row.append(QStandardItem(doc.core_properties.title))
                            self.ExpTableModel.appendRow(row)
                            self.ExpTable.resizeRowsToContents()
                            Experiments[file.split()[1]] = root+"\\"+file
        doc = docx.Document(os.path.realpath(".\demo 1.docx"))
        print(doc.paragraphs[1].text)
        print(Experiments)

    #Open calendar function
    def OpenCal(self):
        print("hello")
        self.DTNew_Cal.show()

    def AddSection(self):
        print("hello")

    #When the Variable Category ComboBox is changed, update the second combobox
    def VarChange(self, s):
        self.DTNew_Var_Comb2.clear()
        print(s)
        var = s
        conn = sqlite3.connect("./Variables/Variables.sqlite3")
        rows = conn.execute("SELECT name FROM {0};".format(var))
        templist = []
        for name in rows:
            templist.append(name[0])
        self.DTNew_Var_Comb2.addItems(templist)
        conn.close()

    def InsVar(self):
        print(self.DTNew_Var_Comb.currentText(),self.DTNew_Var_Comb2.currentText())
        


# Is this file being run directly?
if __name__ == '__main__':
    app = QApplication(sys.argv)

    #
    start = Atlas()
    #

    sys.exit(app.exec_())
